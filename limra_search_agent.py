"""
LIMRA 문서 검색 및 다운로드 에이전트
- 로그인 후 문서/아티클 검색
- PDF 및 문서 파일 다운로드
"""

import asyncio
import nest_asyncio
nest_asyncio.apply()  # 중첩 이벤트 루프 허용
import os
import re
import json
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin, urlparse, quote_plus

from playwright.async_api import async_playwright, Page, Browser, BrowserContext
import requests


class LimraSearchAgent:
    """LIMRA 웹사이트 검색 및 다운로드 에이전트"""

    BASE_URL = "https://www.limra.com"
    LOGIN_URL = "https://www.limra.com/login/"
    SEARCH_URL = "https://www.limra.com/en/search/"
    RESEARCH_URLS = [
        "https://www.limra.com/en/research/",
        "https://www.limra.com/en/research/insurance/",
        "https://www.limra.com/en/research/retirement/",
        "https://www.limra.com/en/research/annuities/",
        "https://www.limra.com/en/research/workplace-benefits/",
    ]

    NEWS_URLS = [
        "https://www.limra.com/en/newsroom/",
        "https://www.limra.com/en/newsroom/news-releases/",
        "https://www.limra.com/en/newsroom/industry-trends/",
    ]

    def __init__(
        self,
        email: str,
        password: str,
        download_folder: str = "./downloads",
        headless: bool = False  # 디버깅을 위해 기본값 False
    ):
        self.email = email
        self.password = password
        self.download_folder = Path(download_folder)
        self.headless = headless
        self.browser: Browser = None
        self.context: BrowserContext = None
        self.page: Page = None
        self.is_logged_in = False

        # 다운로드 폴더 생성
        self.download_folder.mkdir(parents=True, exist_ok=True)

        # 검색 결과 저장
        self.search_results = []

    async def initialize(self):
        """브라우저 초기화 (저장된 세션 자동 로드)"""
        print("[*] 브라우저 초기화 중...")
        self._playwright = await async_playwright().start()

        self.browser = await self._playwright.chromium.launch(
            headless=self.headless,
            args=[
                '--disable-blink-features=AutomationControlled',
                '--disable-pdf-viewer',  # PDF 뷰어 비활성화
                '--disable-popup-blocking',  # 팝업 차단 비활성화
                '--disable-save-password-bubble',  # 비밀번호 저장 팝업 비활성화
                f'--download-default-directory={str(self.download_folder.absolute())}',  # 기본 다운로드 경로
                '--disable-download-notification',  # 다운로드 알림 비활성화
            ],
            downloads_path=str(self.download_folder.absolute())  # Playwright 다운로드 경로
        )

        # 저장된 세션 확인
        session_file = self.download_folder / "limra_session.json"
        context_options = {
            'viewport': {'width': 1920, 'height': 1080},
            'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'accept_downloads': True,
        }

        # 저장된 세션이 있으면 자동으로 로드
        if session_file.exists():
            print(f"[SESSION] 저장된 세션 발견: {session_file}")
            try:
                context_options['storage_state'] = str(session_file)
                print("[SESSION] 세션을 자동으로 로드합니다...")
            except Exception as e:
                print(f"[SESSION] 세션 로드 실패: {e}")

        # PDF 자동 다운로드 설정 (세션 포함)
        self.context = await self.browser.new_context(**context_options)

        # 메인 페이지 생성
        self.page = await self.context.new_page()

        # CDP를 통해 메인 페이지에 자동 다운로드 설정
        self._cdp_session = await self.context.new_cdp_session(self.page)
        await self._cdp_session.send('Page.setDownloadBehavior', {
            'behavior': 'allow',
            'downloadPath': str(self.download_folder.absolute())
        })

        # Browser 레벨 다운로드 설정 추가
        await self._cdp_session.send('Browser.setDownloadBehavior', {
            'behavior': 'allow',
            'downloadPath': str(self.download_folder.absolute()),
            'eventsEnabled': True
        })

        # 자동화 감지 우회
        await self.page.add_init_script("""
            Object.defineProperty(navigator, 'webdriver', {
                get: () => undefined
            });
        """)

        print("[OK] 브라우저 초기화 완료")

    async def load_session(self) -> bool:
        """저장된 세션 쿠키 로드"""
        session_file = self.download_folder / "limra_session.json"

        if session_file.exists():
            try:
                print(f"[SESSION] 저장된 세션 로드 중: {session_file}")
                with open(session_file, 'r') as f:
                    storage_state = json.load(f)

                # 쿠키 복원
                if 'cookies' in storage_state:
                    await self.context.add_cookies(storage_state['cookies'])
                    print(f"[SESSION] {len(storage_state['cookies'])}개 쿠키 로드 완료")

                    # 세션 유효성 확인
                    await self.page.goto(self.BASE_URL)
                    await asyncio.sleep(2)

                    # 로그인 상태 확인 (여러 방법)
                    is_logged_in = await self._check_if_logged_in()

                    if is_logged_in:
                        self.is_logged_in = True
                        print("[SESSION] [OK] 세션이 유효합니다. 로그인 건너뜀")
                        return True
                    else:
                        print("[SESSION] [WARN] 세션이 만료되었습니다. 재로그인 필요")
                        return False
            except Exception as e:
                print(f"[SESSION] 세션 로드 실패: {e}")
                return False
        else:
            print("[SESSION] 저장된 세션이 없습니다")
            return False

    async def _check_if_logged_in(self) -> bool:
        """현재 페이지에서 로그인 상태 확인"""
        try:
            current_url = self.page.url
            content = await self.page.content()
            content_lower = content.lower()

            # 방법 1: URL 확인 - 로그인 페이지가 아니면 로그인 상태일 가능성 높음
            if 'login' in current_url.lower() or 'log-in' in current_url.lower():
                print("[CHECK] URL에 login 포함 - 로그인 안됨")
                return False

            # 방법 2: 로그아웃 버튼/링크 확인 (영어/한국어)
            logout_indicators = [
                'logout', 'log out', 'sign out', 'signout',
                '로그아웃', '로그 아웃', '나가기'
            ]
            for indicator in logout_indicators:
                if indicator in content_lower:
                    print(f"[CHECK] '{indicator}' 발견 - 로그인 상태 확인")
                    return True

            # 방법 3: 로그인 버튼/링크 확인 - 있으면 로그인 안된 상태
            login_indicators = [
                'href="/login"', 'href="/log-in"', 'href="/en/log-in"',
                '>log in<', '>login<', '>sign in<',
                '>로그인<'
            ]
            for indicator in login_indicators:
                if indicator in content_lower:
                    print(f"[CHECK] '{indicator}' 발견 - 로그인 안됨")
                    return False

            # 방법 4: My Account, Profile 등 로그인 후에만 보이는 요소 확인
            account_indicators = [
                'my account', 'myaccount', 'my profile', 'myprofile',
                'welcome,', '내 계정', '프로필'
            ]
            for indicator in account_indicators:
                if indicator in content_lower:
                    print(f"[CHECK] '{indicator}' 발견 - 로그인 상태 확인")
                    return True

            # 방법 5: 쿠키 확인 - 인증 관련 쿠키가 있는지
            cookies = await self.context.cookies()
            auth_cookie_names = ['auth', 'token', 'session', 'user', '.aspnet']
            for cookie in cookies:
                cookie_name_lower = cookie.get('name', '').lower()
                for auth_name in auth_cookie_names:
                    if auth_name in cookie_name_lower:
                        print(f"[CHECK] 인증 쿠키 '{cookie['name']}' 발견")
                        # 쿠키만으로는 확정할 수 없으므로 URL 체크와 함께 판단
                        if 'limra.com' in current_url and 'login' not in current_url.lower():
                            return True

            print("[CHECK] 로그인 상태 불확실 - 로그인 안된 것으로 간주")
            return False

        except Exception as e:
            print(f"[CHECK] 로그인 상태 확인 오류: {e}")
            return False

    async def manual_login_once(self) -> bool:
        """최초 1회 수동 로그인 - 브라우저를 열어두고 충분한 시간 제공"""
        print("\n" + "="*60)
        print("[MANUAL LOGIN] 수동 로그인 모드")
        print("="*60)
        print("브라우저가 열립니다. 다음 작업을 수행하세요:")
        print("1. 이메일과 비밀번호 입력")
        print("2. CAPTCHA 해결")
        print("3. 로그인 완료 후 메인 페이지 확인")
        print("\n120초 대기 후 자동으로 세션을 저장합니다...")
        print("="*60 + "\n")

        try:
            # 로그인 페이지로 이동
            await self.page.goto(self.LOGIN_URL, wait_until='networkidle', timeout=60000)

            # 사용자가 수동으로 로그인할 시간 제공 (2분)
            await asyncio.sleep(120)

            # 로그인 확인
            current_url = self.page.url
            page_content = await self.page.content()

            if any([
                'logout' in page_content.lower(),
                'sign out' in page_content.lower(),
                'limra.com' in current_url and 'login' not in current_url.lower()
            ]):
                print("[SUCCESS] 로그인 성공! 세션을 저장합니다...")
                self.is_logged_in = True

                # 세션 저장
                storage_state = await self.context.storage_state()
                session_path = self.download_folder / "limra_session.json"
                with open(session_path, 'w') as f:
                    json.dump(storage_state, f)
                print(f"[SUCCESS] 세션 저장 완료: {session_path}")
                print("[INFO] 다음번부터는 CAPTCHA 없이 자동 로그인됩니다!\n")

                return True
            else:
                print("[ERROR] 로그인이 완료되지 않았습니다.")
                return False

        except Exception as e:
            print(f"[ERROR] 수동 로그인 실패: {e}")
            return False

    async def login(self) -> bool:
        """LIMRA 웹사이트 로그인 (2단계 로그인 지원)"""
        # 먼저 저장된 세션 시도
        if await self.load_session():
            return True

        # 로그 파일 생성
        log_file = self.download_folder / f"login_debug_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

        def log(msg):
            timestamp = datetime.now().strftime('%H:%M:%S')
            log_msg = f"[{timestamp}] {msg}"
            print(log_msg)
            with open(log_file, 'a', encoding='utf-8') as f:
                f.write(log_msg + '\n')

        log(f"[*] 로그인 시도 시작... ({self.email})")
        log(f"[DEBUG] 로그 파일: {log_file}")

        try:
            # 로그인 페이지로 이동
            log(f"[STEP 1] 로그인 페이지로 이동 중: {self.LOGIN_URL}")
            await self.page.goto(self.LOGIN_URL, wait_until='networkidle', timeout=60000)
            log("[STEP 1] 페이지 로드 완료, 3초 대기 중...")
            await asyncio.sleep(3)

            # 현재 URL 확인 (리다이렉트 될 수 있음)
            current_url = self.page.url
            log(f"[STEP 1] 현재 페이지: {current_url}")

            # 이미 로그인 상태인지 확인 (로그인 페이지로 갔는데 메인으로 리다이렉트된 경우)
            if 'limra.com' in current_url and 'login' not in current_url.lower() and 'account' not in current_url.lower():
                log("[STEP 1] 이미 로그인 상태로 감지됨! 로그인 건너뜀")
                self.is_logged_in = True
                # 세션 저장
                try:
                    storage_state = await self.context.storage_state()
                    session_path = self.download_folder / "limra_session.json"
                    with open(session_path, 'w') as f:
                        json.dump(storage_state, f)
                    log(f"[SESSION] 세션 저장됨: {session_path}")
                except Exception as save_err:
                    log(f"[WARN] 세션 저장 실패: {save_err}")
                return True

            # === 1단계: 이메일 입력 ===
            log("\n[STEP 2] 이메일 입력 단계 시작...")

            # 이메일 입력 필드 찾기
            log("[STEP 2] 이메일 입력 필드 찾는 중...")
            email_selectors = [
                'input[type="email"]',
                'input[name="email"]',
                'input[name="Email"]',
                'input[name="username"]',
                'input[name="Username"]',
                'input[id="email"]',
                'input[id="Email"]',
                '#Email',
                '.email-input',
                'input[placeholder*="email" i]',
                'input[placeholder*="이메일" i]',
                # 로그인 폼 내의 text input만 선택 (검색창 제외)
                'form input[type="text"]',
            ]

            email_input = None
            for i, selector in enumerate(email_selectors, 1):
                try:
                    log(f"[STEP 2] 시도 {i}/{len(email_selectors)}: {selector}")
                    email_input = await self.page.wait_for_selector(selector, timeout=3000)
                    if email_input and await email_input.is_visible():
                        log(f"[OK] 이메일 필드 발견: {selector}")
                        break
                except Exception as e:
                    log(f"[SKIP] {selector} - {type(e).__name__}")
                    continue

            if not email_input:
                # 페이지 소스 저장하여 디버깅
                log("[ERROR] 이메일 필드를 찾을 수 없음! 디버그 정보 저장 중...")
                html = await self.page.content()
                debug_path = self.download_folder / "debug_login_page.html"
                with open(debug_path, 'w', encoding='utf-8') as f:
                    f.write(html)
                log(f"[SAVED] 페이지 소스: {debug_path}")

                # 스크린샷 저장
                screenshot_path = self.download_folder / "debug_login_screenshot.png"
                await self.page.screenshot(path=str(screenshot_path))
                log(f"[SAVED] 스크린샷: {screenshot_path}")
                return False

            # 이메일 입력
            log(f"[STEP 2] 이메일 입력 중: {self.email}")
            await email_input.fill(self.email)
            log("[STEP 2] 이메일 입력 완료, 1초 대기...")
            await asyncio.sleep(1)

            # reCAPTCHA 체크 (선택적, 90초 대기)
            log("[STEP 3] reCAPTCHA 확인 중...")
            captcha_frame = None
            try:
                # reCAPTCHA iframe 찾기
                captcha_frame = await self.page.wait_for_selector('iframe[src*="recaptcha"]', timeout=3000)
                if captcha_frame:
                    log("[!] reCAPTCHA 감지됨! 수동으로 체크박스를 클릭하세요.")
                    log("[!] 90초 대기 중...")

                    # CAPTCHA가 해결될 때까지 대기 (최대 90초)
                    for i in range(90):
                        await asyncio.sleep(1)

                        # 페이지 URL 확인 (CAPTCHA 완료 후 자동 로그인되면 페이지가 이동함)
                        try:
                            current_url = self.page.url
                            if 'limra.com' in current_url and 'login' not in current_url.lower() and 'discovery' not in current_url.lower():
                                log(f"[OK] CAPTCHA 완료 후 자동 로그인됨! URL: {current_url}")
                                break
                        except:
                            pass

                        # CAPTCHA 응답 확인
                        try:
                            captcha_response = await self.page.evaluate('''
                                () => {
                                    const response = document.querySelector('[name="g-recaptcha-response"]');
                                    return response ? response.value : null;
                                }
                            ''')
                            if captcha_response:
                                log(f"[OK] reCAPTCHA 해결됨! ({i+1}초 경과)")
                                await asyncio.sleep(2)  # CAPTCHA 완료 후 페이지 전환 대기
                                break
                        except:
                            pass

                        if i % 10 == 9:
                            log(f"[...] {i+1}초 경과... CAPTCHA를 완료해주세요.")
                    else:
                        log("[WARN] CAPTCHA 대기 시간 초과")

            except Exception as captcha_err:
                log(f"[i] reCAPTCHA 없음: {type(captcha_err).__name__}")

            log("[STEP 3] CAPTCHA 확인 완료, 2초 대기...")
            await asyncio.sleep(2)

            # CAPTCHA 완료 후 페이지 상태 확인 (자동 로그인 감지)
            try:
                current_url = self.page.url
                log(f"[STEP 3.5] CAPTCHA 후 현재 URL: {current_url}")

                # 이미 메인 페이지로 이동했다면 로그인 성공
                if 'limra.com' in current_url and 'login' not in current_url.lower():
                    log("[SUCCESS] CAPTCHA 완료 후 자동으로 로그인되었습니다!")
                    self.is_logged_in = True

                    # 세션 저장
                    try:
                        storage_state = await self.context.storage_state()
                        session_path = self.download_folder / "limra_session.json"
                        with open(session_path, 'w') as f:
                            json.dump(storage_state, f)
                        log(f"[SESSION] 세션 저장됨: {session_path}")
                    except Exception as save_err:
                        log(f"[WARN] 세션 저장 실패: {type(save_err).__name__}")

                    return True
            except Exception as e:
                log(f"[WARN] URL 확인 실패: {type(e).__name__}, 계속 진행...")

            # 페이지가 닫혔는지 확인
            try:
                current_url = self.page.url
                log(f"[STEP 4] 버튼 찾기 전 URL: {current_url}")
            except Exception as e:
                log(f"[ERROR] 페이지가 이미 닫혔습니다: {type(e).__name__}")
                # 새 페이지를 다시 가져오기 시도
                try:
                    pages = self.context.pages
                    if pages:
                        self.page = pages[0]
                        log(f"[RECOVER] 첫 번째 페이지로 복구: {self.page.url}")
                    else:
                        log("[ERROR] 사용 가능한 페이지가 없습니다")
                        return False
                except Exception as e2:
                    log(f"[ERROR] 페이지 복구 실패: {type(e2).__name__}")
                    return False

            # 로그인/다음 버튼 클릭 (1단계)
            log("[STEP 4] 첫 번째 버튼 찾는 중...")
            first_button_selectors = [
                'button[type="submit"]',
                'input[type="submit"]',
                'button:has-text("로그인")',
                'button:has-text("Log In")',
                'button:has-text("Login")',
                'button:has-text("Next")',
                'button:has-text("다음")',
                'button:has-text("Continue")',
                '.btn-primary',
                'button.btn',
            ]

            first_button = None
            for i, selector in enumerate(first_button_selectors, 1):
                try:
                    log(f"[STEP 4] 버튼 시도 {i}/{len(first_button_selectors)}: {selector}")
                    first_button = await self.page.wait_for_selector(selector, timeout=2000)
                    if first_button and await first_button.is_visible():
                        log(f"[OK] 1단계 버튼 발견: {selector}")
                        break
                except Exception as e:
                    log(f"[SKIP] {selector} - {type(e).__name__}")
                    # TargetClosedError면 페이지가 닫혔으므로 복구 시도
                    if 'TargetClosedError' in str(type(e).__name__):
                        log("[WARN] 페이지가 닫혔습니다. 현재 페이지 확인 중...")
                        try:
                            pages = self.context.pages
                            if pages:
                                self.page = pages[0]
                                log(f"[RECOVER] 페이지 복구됨: {self.page.url}")
                                # 이미 로그인되었는지 확인
                                if 'limra.com' in self.page.url and 'login' not in self.page.url.lower():
                                    log("[CHECK] 로그인 페이지가 아닙니다. 로그인 상태 확인 중...")
                                    break
                        except:
                            pass
                    continue

            # 버튼이 없으면 이미 로그인되었을 수 있음
            if not first_button:
                log("[STEP 4] 버튼을 찾지 못함. 이미 로그인되었는지 확인 중...")
                try:
                    current_url = self.page.url
                    if 'limra.com' in current_url and 'login' not in current_url.lower():
                        log("[STEP 4] 로그인 페이지가 아닙니다. STEP 5로 건너뜁니다.")
                        # 비밀번호 단계로 직접 이동
                    else:
                        log("[STEP 4] 버튼 없음, Enter 키 시도...")
                        try:
                            async with self.page.expect_navigation(timeout=15000, wait_until='networkidle'):
                                await email_input.press('Enter')
                            log("[STEP 4] Enter 키 완료 (페이지 이동됨)")
                        except Exception as nav_err:
                            log(f"[i] Enter 키 완료 ({type(nav_err).__name__})")
                except Exception as e:
                    log(f"[WARN] URL 확인 중 오류: {type(e).__name__}")

            # 첫 번째 버튼 클릭 (네비게이션 발생 대비)
            elif first_button:
                log("[STEP 4] 버튼 클릭 시도 중...")
                try:
                    # 네비게이션이 발생할 수 있으므로 expect_navigation 사용
                    log("[STEP 4] 버튼 클릭 실행...")
                    try:
                        async with self.page.expect_navigation(timeout=15000, wait_until='networkidle'):
                            await first_button.click()
                        log("[STEP 4] 버튼 클릭 완료 (페이지 이동됨)")
                    except Exception as nav_err:
                        # 네비게이션이 없을 수도 있음 - 무시
                        log(f"[i] 1단계 클릭 완료 ({type(nav_err).__name__})")
                except Exception as click_err:
                    log(f"[WARN] 1단계 버튼 클릭 중 오류: {click_err}")
                    # 클릭 오류 시에도 계속 진행

            log("[...] 다음 단계 대기 중... (4초)")
            await asyncio.sleep(4)

            try:
                log("[STEP 4] 페이지 로드 대기 중...")
                await self.page.wait_for_load_state('networkidle', timeout=30000)
                log("[STEP 4] 페이지 로드 완료")
            except Exception as e:
                log(f"[STEP 4] 페이지 로드 타임아웃 ({type(e).__name__}), 계속 진행...")

            # === 2단계: 비밀번호 입력 ===
            log("\n[STEP 5] 비밀번호 입력 단계 시작...")

            # 페이지 복구 확인
            try:
                current_url = self.page.url
                log(f"[STEP 5] 현재 URL: {current_url}")
            except Exception as e:
                log(f"[WARN] 페이지 URL 접근 실패: {type(e).__name__}")
                try:
                    pages = self.context.pages
                    if pages:
                        self.page = pages[0]
                        log(f"[RECOVER] 페이지 복구됨: {self.page.url}")
                    else:
                        log("[ERROR] 사용 가능한 페이지 없음")
                        return False
                except Exception as e2:
                    log(f"[ERROR] 페이지 복구 실패: {type(e2).__name__}")
                    return False

            # 비밀번호 필드 찾기
            log("[STEP 5] 비밀번호 필드 찾는 중...")
            password_selectors = [
                'input[type="password"]',
                'input[name="password"]',
                'input[name="Password"]',
                '#Password',
                '#password',
                'input[placeholder*="password" i]',
                'input[placeholder*="비밀번호" i]',
            ]

            password_input = None
            for i, selector in enumerate(password_selectors, 1):
                try:
                    log(f"[STEP 5] 비밀번호 시도 {i}/{len(password_selectors)}: {selector}")
                    password_input = await self.page.wait_for_selector(selector, timeout=5000)
                    if password_input and await password_input.is_visible():
                        log(f"[OK] 비밀번호 필드 발견: {selector}")
                        break
                except Exception as e:
                    log(f"[SKIP] {selector} - {type(e).__name__}")
                    # TargetClosedError면 페이지 복구 시도
                    if 'TargetClosedError' in str(type(e).__name__):
                        log("[WARN] 페이지가 닫혔습니다. 복구 시도 중...")
                        try:
                            pages = self.context.pages
                            if pages:
                                self.page = pages[0]
                                log(f"[RECOVER] 페이지 복구됨: {self.page.url}")
                                break
                        except:
                            pass
                    continue

            if not password_input:
                log("[WARN] 비밀번호 필드를 찾을 수 없습니다.")

                # 이미 로그인 된 상태일 수 있음
                log("[CHECK] 로그인 상태 확인 중...")
                try:
                    current_url = self.page.url
                    page_content = await self.page.content()

                    if 'logout' in page_content.lower() or 'sign out' in page_content.lower():
                        log("[SUCCESS] 이미 로그인된 상태입니다!")
                        self.is_logged_in = True

                        # 세션 저장
                        storage_state = await self.context.storage_state()
                        session_path = self.download_folder / "limra_session.json"
                        with open(session_path, 'w') as f:
                            json.dump(storage_state, f)
                        log(f"[SESSION] 세션 저장됨: {session_path}")

                        return True

                    # 로그인 페이지가 아니면 로그인 성공으로 간주
                    if 'limra.com' in current_url and 'login' not in current_url.lower():
                        log("[SUCCESS] 로그인 페이지가 아닙니다. 로그인 성공으로 간주합니다.")
                        self.is_logged_in = True

                        # 세션 저장
                        storage_state = await self.context.storage_state()
                        session_path = self.download_folder / "limra_session.json"
                        with open(session_path, 'w') as f:
                            json.dump(storage_state, f)
                        log(f"[SESSION] 세션 저장됨: {session_path}")

                        return True

                except Exception as e:
                    log(f"[ERROR] 로그인 상태 확인 실패: {type(e).__name__}")

                # 스크린샷 저장 시도
                try:
                    screenshot_path = self.download_folder / "debug_password_step.png"
                    await self.page.screenshot(path=str(screenshot_path))
                    log(f"[SAVED] 스크린샷: {screenshot_path}")
                except:
                    log("[WARN] 스크린샷 저장 실패")

                log("[ERROR] 비밀번호 필드 없음, 로그인 실패")
                return False

            # 비밀번호 입력
            log("[STEP 5] 비밀번호 입력 중...")
            await password_input.fill(self.password)
            log("[STEP 5] 비밀번호 입력 완료, 1초 대기...")
            await asyncio.sleep(1)

            # 로그인 버튼 클릭 (2단계)
            log("[STEP 6] 로그인 버튼 찾는 중...")
            login_button_selectors = [
                'button[type="submit"]',
                'input[type="submit"]',
                'button:has-text("로그인")',
                'button:has-text("Log In")',
                'button:has-text("Login")',
                'button:has-text("Sign In")',
                '.login-button',
                '#loginButton',
                'button.btn-primary',
                'button.btn',
            ]

            login_button = None
            for i, selector in enumerate(login_button_selectors, 1):
                try:
                    log(f"[STEP 6] 로그인 버튼 시도 {i}/{len(login_button_selectors)}: {selector}")
                    login_button = await self.page.wait_for_selector(selector, timeout=2000)
                    if login_button and await login_button.is_visible():
                        log(f"[OK] 로그인 버튼 발견: {selector}")
                        break
                except Exception as e:
                    log(f"[SKIP] {selector} - {type(e).__name__}")
                    continue

            # 로그인 버튼 클릭 (네비게이션 발생 시 예외 무시)
            log("[STEP 6] 로그인 버튼 클릭 시도...")
            try:
                if login_button:
                    log("[STEP 6] 로그인 버튼 클릭 실행...")
                    # Promise.all 패턴으로 클릭과 네비게이션 동시 처리
                    async with self.page.expect_navigation(timeout=30000, wait_until='networkidle'):
                        await login_button.click()
                    log("[STEP 6] 로그인 버튼 클릭 완료 (페이지 이동됨)")
                else:
                    log("[STEP 6] 로그인 버튼 없음, Enter 키 시도...")
                    async with self.page.expect_navigation(timeout=30000, wait_until='networkidle'):
                        await password_input.press('Enter')
                    log("[STEP 6] Enter 키 완료 (페이지 이동됨)")
            except Exception as nav_error:
                # 네비게이션 중 요소 분리 오류는 무시 (실제로 로그인 성공했을 가능성)
                log(f"[...] 페이지 네비게이션 중... ({type(nav_error).__name__})")
                await asyncio.sleep(3)

            # 로그인 완료 대기
            log("[STEP 6] 로그인 처리 중... (3초 대기)")
            await asyncio.sleep(3)

            try:
                log("[STEP 7] 페이지 로드 대기 중...")
                await self.page.wait_for_load_state('networkidle', timeout=15000)
                log("[STEP 7] 페이지 로드 완료")
            except Exception as e:
                log(f"[STEP 7] 페이지 로드 타임아웃 ({type(e).__name__}), 계속 진행...")

            # 로그인 성공 확인
            current_url = self.page.url
            log(f"[STEP 7] 로그인 후 URL: {current_url}")

            log("[STEP 7] 페이지 콘텐츠 분석 중...")
            page_content = await self.page.content()

            # 로그인 성공 지표 확인
            log("[STEP 7] 로그인 성공 지표 확인 중...")
            success_indicators = [
                'logout' in page_content.lower(),
                'sign out' in page_content.lower(),
                'my account' in page_content.lower(),
                'my limra' in page_content.lower(),
                'welcome' in page_content.lower(),
                '로그아웃' in page_content,
                '회원' in page_content,
                'www.limra.com' in current_url and 'login' not in current_url.lower(),
            ]

            log(f"[STEP 7] 성공 지표 확인 결과: {success_indicators}")

            if any(success_indicators):
                self.is_logged_in = True
                log("[SUCCESS] 로그인 성공!")

                # 전체 세션 상태 저장 (쿠키 + localStorage)
                storage_state = await self.context.storage_state()
                session_path = self.download_folder / "limra_session.json"
                with open(session_path, 'w') as f:
                    json.dump(storage_state, f)
                log(f"[SESSION] 세션 저장됨: {session_path}")

                return True
            else:
                log("[WARN] 로그인 지표를 찾지 못함, 메인 페이지에서 재확인...")
                # 메인 페이지로 이동해서 다시 확인
                log("[STEP 8] 메인 페이지로 이동 중...")
                await self.page.goto(self.BASE_URL, wait_until='networkidle', timeout=30000)
                await asyncio.sleep(2)

                page_content = await self.page.content()
                current_url = self.page.url
                log(f"[STEP 8] 메인 페이지 URL: {current_url}")

                if any([
                    'logout' in page_content.lower(),
                    'sign out' in page_content.lower(),
                    'my limra' in page_content.lower(),
                    '로그아웃' in page_content,
                    'limra.com' in current_url and 'login' not in current_url.lower(),
                ]):
                    self.is_logged_in = True
                    log("[SUCCESS] 로그인 성공! (메인 페이지에서 확인)")

                    # 세션 저장
                    storage_state = await self.context.storage_state()
                    session_path = self.download_folder / "limra_session.json"
                    with open(session_path, 'w') as f:
                        json.dump(storage_state, f)
                    log(f"[SESSION] 세션 저장됨: {session_path}")

                    return True

                # 추가 확인: 로그인 페이지가 아니면 로그인 성공으로 간주
                if 'www.limra.com' in current_url and 'login' not in current_url.lower():
                    self.is_logged_in = True
                    log("[SUCCESS] 로그인 성공! (URL 확인)")

                    # 세션 저장
                    storage_state = await self.context.storage_state()
                    session_path = self.download_folder / "limra_session.json"
                    with open(session_path, 'w') as f:
                        json.dump(storage_state, f)
                    log(f"[SESSION] 세션 저장됨: {session_path}")

                    return True

                log("[ERROR] 로그인 실패 - 자격 증명을 확인해주세요")
                screenshot_path = self.download_folder / "login_failed_screenshot.png"
                await self.page.screenshot(path=str(screenshot_path))
                log(f"[SAVED] 실패 스크린샷: {screenshot_path}")
                return False

        except Exception as e:
            # 오류가 발생해도 로그인 상태 확인
            log(f"[EXCEPTION] 예외 발생: {type(e).__name__} - {str(e)}")
            log("[*] 로그인 상태 재확인 중...")

            try:
                log("[EXCEPTION] 2초 대기 후 메인 페이지로 이동...")
                await asyncio.sleep(2)
                await self.page.goto(self.BASE_URL, wait_until='networkidle', timeout=30000)
                page_content = await self.page.content()
                log(f"[EXCEPTION] 현재 URL: {self.page.url}")

                if any([
                    'logout' in page_content.lower(),
                    'sign out' in page_content.lower(),
                    'my limra' in page_content.lower(),
                ]):
                    self.is_logged_in = True
                    log("[SUCCESS] 예외 발생 후 로그인 성공 확인!")

                    # 세션 저장
                    storage_state = await self.context.storage_state()
                    session_path = self.download_folder / "limra_session.json"
                    with open(session_path, 'w') as f:
                        json.dump(storage_state, f)
                    log(f"[SESSION] 세션 저장됨: {session_path}")

                    return True
            except Exception as e2:
                log(f"[EXCEPTION] 재확인 중 오류: {type(e2).__name__}")

            log("[ERROR] 최종 로그인 실패")
            screenshot_path = self.download_folder / "login_error_screenshot.png"
            try:
                await self.page.screenshot(path=str(screenshot_path))
                log(f"[SAVED] 오류 스크린샷: {screenshot_path}")
            except Exception as e3:
                log(f"[ERROR] 스크린샷 저장 실패: {type(e3).__name__}")
            return False

    async def search_documents(self, query: str, max_results: int = 50) -> list:
        """문서 검색"""
        print(f"\n[SEARCH] 검색 중: '{query}'")

        results = []

        try:
            # 검색 페이지로 이동
            search_url = f"{self.SEARCH_URL}?q={quote_plus(query)}"
            await self.page.goto(search_url, wait_until='networkidle', timeout=60000)
            await asyncio.sleep(3)

            # 검색 결과 파싱
            results = await self._parse_search_results(max_results)

            # 추가 연구 섹션에서도 검색
            for research_url in self.RESEARCH_URLS:
                try:
                    await self.page.goto(research_url, wait_until='networkidle', timeout=30000)
                    await asyncio.sleep(2)

                    # 페이지 내 검색 기능이 있으면 사용
                    search_input = await self.page.query_selector('input[type="search"], input[name="q"], .search-input')
                    if search_input:
                        await search_input.fill(query)
                        await search_input.press('Enter')
                        await asyncio.sleep(3)

                        additional_results = await self._parse_search_results(max_results - len(results))
                        results.extend(additional_results)

                except Exception as e:
                    print(f"[WARN] {research_url} 검색 중 오류: {e}")
                    continue

                if len(results) >= max_results:
                    break

        except Exception as e:
            print(f"[ERROR] 검색 중 오류 발생: {e}")

        # 중복 제거
        seen_urls = set()
        unique_results = []
        for result in results:
            if result['url'] not in seen_urls:
                seen_urls.add(result['url'])
                unique_results.append(result)

        self.search_results = unique_results[:max_results]
        print(f"[OK] 총 {len(self.search_results)}개 결과 발견")

        return self.search_results

    async def _parse_search_results(self, max_results: int) -> list:
        """검색 결과 페이지 파싱"""
        results = []

        # 다양한 결과 컨테이너 셀렉터
        result_selectors = [
            '.search-result',
            '.search-results-item',
            '.result-item',
            'article',
            '.card',
            '.list-item',
            '[class*="result"]',
            '[class*="article"]',
        ]

        for selector in result_selectors:
            items = await self.page.query_selector_all(selector)
            if items:
                print(f"[LIST] {len(items)}개 항목 발견 ({selector})")

                for item in items[:max_results]:
                    try:
                        result = await self._extract_result_info(item)
                        if result:
                            results.append(result)
                    except Exception as e:
                        continue

                if results:
                    break

        # 직접 링크 검색 (PDF, 문서 등)
        all_links = await self.page.query_selector_all('a[href]')
        for link in all_links:
            try:
                href = await link.get_attribute('href')
                text = await link.inner_text()

                if href and self._is_document_link(href):
                    full_url = urljoin(self.BASE_URL, href)
                    results.append({
                        'title': text.strip() or 'Untitled Document',
                        'url': full_url,
                        'type': self._get_document_type(href),
                        'description': ''
                    })
            except:
                continue

        return results

    async def _extract_result_info(self, element) -> dict:
        """개별 검색 결과 정보 추출"""
        try:
            # 제목 추출
            title_el = await element.query_selector('h1, h2, h3, h4, .title, [class*="title"]')
            title = await title_el.inner_text() if title_el else ''

            # 링크 추출
            link_el = await element.query_selector('a[href]')
            url = await link_el.get_attribute('href') if link_el else ''

            if not url:
                return None

            url = urljoin(self.BASE_URL, url)

            # 설명 추출
            desc_el = await element.query_selector('p, .description, .summary, [class*="desc"]')
            description = await desc_el.inner_text() if desc_el else ''

            # 문서 타입 판별
            doc_type = self._get_document_type(url)

            return {
                'title': title.strip(),
                'url': url,
                'type': doc_type,
                'description': description.strip()[:200]
            }
        except:
            return None

    def _is_document_link(self, url: str) -> bool:
        """문서 링크인지 확인"""
        doc_extensions = ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx']
        url_lower = url.lower()
        return any(ext in url_lower for ext in doc_extensions)

    def _get_document_type(self, url: str) -> str:
        """URL에서 문서 타입 추출"""
        url_lower = url.lower()
        if '.pdf' in url_lower:
            return 'PDF'
        elif '.doc' in url_lower:
            return 'Word'
        elif '.xls' in url_lower:
            return 'Excel'
        elif '.ppt' in url_lower:
            return 'PowerPoint'
        else:
            return 'Article'

    async def _dismiss_cookie_banner(self):
        """쿠키 동의 배너 제거/숨기기"""
        try:
            # 쿠키 배너 닫기 버튼 클릭 시도
            cookie_button_selectors = [
                'button:has-text("Accept All Cookies")',
                'button:has-text("Accept All")',
                'button:has-text("Accept")',
                'button:has-text("동의")',
                'button:has-text("모두 수락")',
                '[id*="cookie"] button',
                '[class*="cookie"] button',
                '.cookie-banner button',
                '#onetrust-accept-btn-handler',
                '.onetrust-close-btn-handler',
            ]

            for selector in cookie_button_selectors:
                try:
                    button = await self.page.query_selector(selector)
                    if button and await button.is_visible():
                        await button.click()
                        await asyncio.sleep(1)
                        return
                except:
                    continue

            # 버튼을 못 찾으면 JavaScript로 쿠키 배너 요소 숨기기
            await self.page.evaluate("""
                () => {
                    // 쿠키 관련 요소 숨기기
                    const selectors = [
                        '[id*="cookie"]',
                        '[class*="cookie"]',
                        '[id*="Cookie"]',
                        '[class*="Cookie"]',
                        '[id*="consent"]',
                        '[class*="consent"]',
                        '[id*="onetrust"]',
                        '[class*="onetrust"]',
                        '[id*="gdpr"]',
                        '[class*="gdpr"]',
                        '.cc-banner',
                        '.cookie-notice',
                        '.cookie-popup',
                        '#cookie-law-info-bar',
                    ];

                    selectors.forEach(selector => {
                        document.querySelectorAll(selector).forEach(el => {
                            el.style.display = 'none';
                            el.style.visibility = 'hidden';
                        });
                    });

                    // position:fixed 요소 중 하단에 위치한 배너 숨기기
                    document.querySelectorAll('*').forEach(el => {
                        const style = window.getComputedStyle(el);
                        if (style.position === 'fixed' &&
                            (style.bottom === '0px' || parseInt(style.bottom) < 100)) {
                            const text = el.innerText.toLowerCase();
                            if (text.includes('cookie') || text.includes('accept') ||
                                text.includes('consent') || text.includes('privacy')) {
                                el.style.display = 'none';
                            }
                        }
                    });
                }
            """)
            await asyncio.sleep(0.5)

        except Exception as e:
            # 쿠키 배너 제거 실패해도 계속 진행
            pass

    async def _dismiss_modal_popup(self):
        """모달 팝업 닫기 (LimraModal 등)"""
        try:
            # LIMRA 사이트의 모달 팝업 닫기
            await self.page.evaluate("""
                () => {
                    // LimraModal 및 일반 모달 닫기
                    const modalSelectors = [
                        '#LimraModal',
                        '.modal',
                        '[class*="modal"]',
                        '.mod-wrapper',
                        '.autofill__panel',
                        '[role="dialog"]',
                        '.popup',
                        '.overlay'
                    ];

                    modalSelectors.forEach(selector => {
                        document.querySelectorAll(selector).forEach(el => {
                            el.style.display = 'none';
                            el.style.visibility = 'hidden';
                            el.style.opacity = '0';
                            el.style.pointerEvents = 'none';
                        });
                    });

                    // 닫기 버튼 클릭 시도
                    const closeButtons = document.querySelectorAll(
                        '.close, .modal-close, [aria-label="Close"], .btn-close, ' +
                        'button[class*="close"], .mod-close, .autofill__close'
                    );
                    closeButtons.forEach(btn => {
                        try { btn.click(); } catch(e) {}
                    });

                    // body overflow 복원 (모달이 스크롤 차단한 경우)
                    document.body.style.overflow = 'auto';
                    document.documentElement.style.overflow = 'auto';
                }
            """)
            await asyncio.sleep(0.3)

        except Exception as e:
            # 모달 제거 실패해도 계속 진행
            pass

    async def browse_research_with_filter(
        self,
        keywords: list = None,
        start_year: int = None,
        end_year: int = None,
        auto_download: bool = False
    ) -> list:
        """
        주제 키워드와 연도로 필터링하여 연구 문서 탐색 및 다운로드

        Args:
            keywords: 검색할 주제 키워드 리스트 (예: ["Recruiting", "Retention"])
            start_year: 시작 연도 (예: 2023)
            end_year: 종료 연도 (예: 2024)
            auto_download: True이면 필터링된 문서 자동 다운로드

        Returns:
            필터링된 문서 목록
        """
        print("\n" + "="*60)
        print("[DOC] Research 섹션 탐색 및 필터링")
        print("="*60)

        if keywords:
            print(f"[SEARCH] 주제 키워드: {', '.join(keywords)}")
        if start_year and end_year:
            print(f"[DATE] 기간: {start_year} ~ {end_year}")
        elif start_year:
            print(f"[DATE] 기간: {start_year} 이후")
        elif end_year:
            print(f"[DATE] 기간: {end_year} 이전")
        print("-"*60)

        # 1단계: 모든 문서 수집 (키워드 전달해서 필터링)
        all_documents = await self.browse_research_section(keywords=keywords)

        # 2단계: 각 문서 페이지에서 날짜 정보 수집 (필요한 경우)
        if start_year or end_year:
            print(f"\n[DATE] 문서 날짜 정보 수집 중... ({len(all_documents)}개)")
            all_documents = await self._collect_document_dates(all_documents)

        # 3단계: 필터링
        filtered_docs = self._filter_documents(all_documents, keywords, start_year, end_year)

        print(f"\n[OK] 필터링 결과: {len(filtered_docs)}개 문서")
        print("-"*60)

        for i, doc in enumerate(filtered_docs[:20], 1):
            year_str = f" ({doc.get('year', '날짜 미상')})" if doc.get('year') else ""
            print(f"{i}. [{doc['type']}] {doc['title'][:60]}{year_str}")

        if len(filtered_docs) > 20:
            print(f"   ... 외 {len(filtered_docs) - 20}개")

        # 4단계: 자동 다운로드
        if auto_download and filtered_docs:
            print(f"\n[DL] {len(filtered_docs)}개 문서 다운로드 시작...")
            self.search_results = filtered_docs
            downloaded = await self.download_all_results()
            print(f"[OK] {len(downloaded)}개 파일 다운로드 완료")

        # 결과 저장
        self.search_results = filtered_docs
        return filtered_docs

    async def _collect_document_dates(self, documents: list) -> list:
        """각 문서 페이지를 방문하여 날짜 정보 수집"""
        updated_docs = []

        for i, doc in enumerate(documents):
            try:
                # 진행 상황 표시
                if (i + 1) % 10 == 0:
                    print(f"  [FILE] {i + 1}/{len(documents)} 문서 확인 중...")

                # 이미 연도 정보가 있으면 스킵
                if doc.get('year'):
                    updated_docs.append(doc)
                    continue

                # 문서 페이지 방문
                await self.page.goto(doc['url'], wait_until='networkidle', timeout=30000)
                await asyncio.sleep(1)

                # 날짜 정보 추출
                year = await self._extract_year_from_page()
                doc['year'] = year
                updated_docs.append(doc)

            except Exception as e:
                doc['year'] = None
                updated_docs.append(doc)
                continue

        return updated_docs

    async def _extract_year_from_page(self) -> int:
        """페이지에서 연도 정보 추출"""
        try:
            # JavaScript로 페이지에서 날짜 정보 찾기
            year = await self.page.evaluate("""
                () => {
                    // 다양한 날짜 패턴 셀렉터
                    const dateSelectors = [
                        '.date', '.publish-date', '.published-date',
                        '[class*="date"]', '[class*="Date"]',
                        'time', '[datetime]',
                        '.meta', '.article-meta', '.post-meta',
                        '.byline', '.info'
                    ];

                    // 연도 패턴 (2020-2029)
                    const yearPattern = /\b(202[0-9]|201[0-9])\b/;

                    // 날짜 셀렉터에서 연도 찾기
                    for (const selector of dateSelectors) {
                        const elements = document.querySelectorAll(selector);
                        for (const el of elements) {
                            const text = el.innerText || el.getAttribute('datetime') || '';
                            const match = text.match(yearPattern);
                            if (match) {
                                return parseInt(match[1]);
                            }
                        }
                    }

                    // 페이지 전체에서 연도 패턴 찾기 (메타 태그 등)
                    const metaTags = document.querySelectorAll('meta[content]');
                    for (const meta of metaTags) {
                        const content = meta.getAttribute('content') || '';
                        const match = content.match(yearPattern);
                        if (match) {
                            return parseInt(match[1]);
                        }
                    }

                    return null;
                }
            """)
            return year
        except:
            return None

    def _filter_documents(
        self,
        documents: list,
        keywords: list = None,
        start_year: int = None,
        end_year: int = None
    ) -> list:
        """문서 목록 필터링 - 키워드 매칭 없으면 전체 반환"""
        filtered = []

        for doc in documents:
            # 키워드 필터링 (제목 + 설명 + 타입 + URL에서 검색)
            if keywords:
                title_lower = doc.get('title', '').lower()
                desc_lower = doc.get('description', '').lower()
                type_lower = doc.get('type', '').lower()
                url_lower = doc.get('url', '').lower()
                search_text = f"{title_lower} {desc_lower} {type_lower} {url_lower}"

                keyword_match = any(kw.lower() in search_text for kw in keywords)
                if not keyword_match:
                    continue

            # 연도 필터링
            doc_year = doc.get('year')
            if doc_year:
                if start_year and doc_year < start_year:
                    continue
                if end_year and doc_year > end_year:
                    continue
            # 연도 정보가 없는 문서는 포함 (최신 문서일 가능성)

            filtered.append(doc)

        # 키워드 필터링 후 결과가 없으면 연도 필터만 적용
        if keywords and not filtered and documents:
            print(f"[WARN] 키워드 매칭 결과 없음. 연도 필터만 적용합니다.")
            for doc in documents:
                doc_year = doc.get('year')
                if doc_year:
                    if start_year and doc_year < start_year:
                        continue
                    if end_year and doc_year > end_year:
                        continue
                filtered.append(doc)

        return filtered

    async def browse_research_section(
        self,
        keywords: list = None,
        start_year: int = None,
        end_year: int = None
    ) -> list:
        """
        연구 섹션 탐색하여 실제 문서/보고서 목록 수집 (간소화 버전)

        Args:
            keywords: 검색 키워드 리스트 (옵션)
            start_year: 시작 연도 (옵션)
            end_year: 종료 연도 (옵션)

        Returns:
            문서 목록
        """
        print("\n[DOC] 연구 섹션 탐색 중 (간소화 방식)...")
        if keywords:
            print(f"[KEYWORDS] {', '.join(keywords)}")

        all_documents = []

        for research_url in self.RESEARCH_URLS:
            try:
                print(f"  → {research_url}")
                await self.page.goto(research_url, wait_until='domcontentloaded', timeout=20000)
                await asyncio.sleep(1)

                # 간소화: 모든 링크를 수집하고 PDF/문서 링크만 필터링
                print(f"    [LINK] 모든 링크 수집 중...")
                links = await self.page.query_selector_all('a[href]')
                print(f"    [LINK] 총 {len(links)}개 링크 발견")

                extracted = 0
                skipped = 0

                for link in links:
                    try:
                        href = await link.get_attribute('href')
                        text = (await link.inner_text()).strip()

                        if not href or not text:
                            skipped += 1
                            continue

                        full_url = urljoin(self.BASE_URL, href)
                        url_lower = full_url.lower()

                        # PDF 직접 링크는 무조건 포함
                        if '.pdf' in url_lower:
                            all_documents.append({
                                'title': text or 'PDF Document',
                                'url': full_url,
                                'type': 'PDF',
                                'description': ''
                            })
                            extracted += 1
                            continue

                        # 제외 패턴
                        skip_patterns = [
                            '?epslanguage=', '/login', '/search', '#', 'javascript:',
                            '/en/research/?', '/en/research/insurance/?',
                            '/en/research/retirement/?', '/en/research/annuities/?',
                            '/en/research/workplace-benefits/?'
                        ]
                        if any(p in full_url for p in skip_patterns):
                            skipped += 1
                            continue

                        # 문서 페이지: URL 깊이 3 이상 + 제목 15자 이상
                        path_segments = [s for s in urlparse(full_url).path.split('/') if s]
                        if len(path_segments) >= 3 and len(text) >= 15:
                            all_documents.append({
                                'title': text,
                                'url': full_url,
                                'type': 'Article',
                                'description': ''
                            })
                            extracted += 1
                        else:
                            skipped += 1

                    except:
                        skipped += 1
                        continue

                print(f"    [LINK] 추출: {extracted}개, 제외: {skipped}개")
                print(f"    [TOTAL] 누적 문서: {len(all_documents)}개")

            except Exception as e:
                print(f"  [WARN] 오류: {e}")
                continue

        print(f"\n[DEBUG] 전체 수집된 문서 수: {len(all_documents)}개 (중복 포함)")

        # 중복 제거
        seen = set()
        unique_docs = []
        for doc in all_documents:
            if doc['url'] not in seen and doc['title'] not in seen:
                seen.add(doc['url'])
                seen.add(doc['title'])
                unique_docs.append(doc)

        print(f"[DEBUG] 중복 제거 후: {len(unique_docs)}개")

        # 샘플 문서 제목 출력 (디버깅용)
        print(f"[DEBUG] 샘플 문서 제목 (처음 10개):")
        for i, doc in enumerate(unique_docs[:10], 1):
            print(f"    {i}. {doc['title'][:80]}")

        # 키워드 필터링
        if keywords:
            print(f"[FILTER] 키워드 필터링 중: {', '.join(keywords)}")
            filtered_docs = []
            for doc in unique_docs:
                title_lower = doc['title'].lower()
                desc_lower = doc.get('description', '').lower()
                url_lower = doc.get('url', '').lower()

                # 키워드 중 하나라도 제목, 설명, 또는 URL에 포함되면 선택
                for keyword in keywords:
                    keyword_lower = keyword.lower()
                    if keyword_lower in title_lower or keyword_lower in desc_lower or keyword_lower in url_lower:
                        filtered_docs.append(doc)
                        print(f"    [MATCH] {doc['title'][:60]}")
                        break

            print(f"[FILTER] 키워드 필터링 후: {len(filtered_docs)}개")

            # 필터링 결과가 없으면 전체 문서 반환 (경고 출력)
            if not filtered_docs:
                print(f"[WARN] 키워드 '{', '.join(keywords)}'와 일치하는 문서가 없습니다.")
                print(f"[WARN] 전체 문서 {len(unique_docs)}개를 반환합니다.")
                filtered_docs = unique_docs

            unique_docs = filtered_docs

        print(f"[OK] 총 {len(unique_docs)}개 실제 문서 발견")
        return unique_docs

    async def download_document(self, url: str, filename: str = None) -> str:
        """문서 다운로드 - 실제 파일 다운로드 우선"""
        try:
            print(f"[DL] 다운로드 중: {url}")

            # 파일명 생성
            if not filename:
                parsed = urlparse(url)
                filename = os.path.basename(parsed.path) or 'document'
                # 파일명 정리
                filename = re.sub(r'[<>:"/\\|?*]', '_', filename)

            # 확장자가 없으면 추가
            if '.' not in filename:
                filename += '.pdf'

            filepath = self.download_folder / filename

            # 중복 파일명 처리
            counter = 1
            original_filepath = filepath
            while filepath.exists():
                stem = original_filepath.stem
                suffix = original_filepath.suffix
                filepath = self.download_folder / f"{stem}_{counter}{suffix}"
                counter += 1

            # 페이지 방문
            await self.page.goto(url, wait_until='networkidle', timeout=60000)
            await asyncio.sleep(2)

            # 쿠키 배너 및 모달 팝업 제거
            await self._dismiss_cookie_banner()
            await self._dismiss_modal_popup()

            # 페이지에서 다운로드 가능한 파일 링크/버튼 찾기
            download_element = await self._find_download_element()

            if download_element:
                # 실제 파일 다운로드 - 클릭 방식
                try:
                    print(f"  [LINK] 다운로드 링크 발견, 클릭 중...")

                    # 모달 다시 확인 후 닫기
                    await self._dismiss_modal_popup()

                    # force 옵션으로 클릭 시도
                    async with self.page.expect_download(timeout=60000) as download_info:
                        await download_element.click(force=True)

                    download = await download_info.value

                    # 다운로드 완료 대기
                    await download.path()

                    # 원본 파일명 사용
                    suggested_filename = download.suggested_filename
                    if suggested_filename:
                        # 확장자 유지
                        ext = Path(suggested_filename).suffix
                        if ext:
                            filepath = filepath.with_suffix(ext)

                    await download.save_as(str(filepath))
                    print(f"[OK] 실제 파일 저장됨: {filepath}")
                    return str(filepath)

                except Exception as e:
                    print(f"[WARN] 클릭 다운로드 실패: {e}")

            # 직접 PDF 링크로 다운로드 시도
            pdf_link = await self._find_pdf_url()
            if pdf_link:
                try:
                    print(f"  [LINK] PDF URL 발견: {pdf_link[:50]}...")

                    # 방법 1: JavaScript 클릭으로 다운로드 시도
                    try:
                        async with self.page.expect_download(timeout=30000) as download_info:
                            # JavaScript로 다운로드 트리거
                            await self.page.evaluate(f'''
                                () => {{
                                    const a = document.createElement('a');
                                    a.href = "{pdf_link}";
                                    a.download = "";
                                    a.style.display = "none";
                                    document.body.appendChild(a);
                                    a.click();
                                    document.body.removeChild(a);
                                }}
                            ''')

                        download = await download_info.value
                        suggested_filename = download.suggested_filename
                        if suggested_filename:
                            ext = Path(suggested_filename).suffix
                            if ext:
                                filepath = filepath.with_suffix(ext)

                        await download.save_as(str(filepath))

                        # 파일 크기 검증
                        if filepath.exists() and filepath.stat().st_size > 0:
                            print(f"[OK] 실제 파일 저장됨: {filepath}")
                            return str(filepath)
                        else:
                            print(f"[WARN] 파일이 비어있음, 다른 방법 시도")
                            if filepath.exists():
                                filepath.unlink()  # 빈 파일 삭제
                    except Exception as js_err:
                        print(f"[WARN] JS 다운로드 실패: {js_err}")

                    # 방법 2: Playwright로 PDF URL 직접 navigate하여 다운로드
                    print(f"  [LINK] Playwright로 PDF 직접 다운로드 시도...")
                    try:
                        # 새 페이지에서 PDF URL로 직접 이동하여 다운로드
                        async with self.page.expect_download(timeout=60000) as download_info:
                            # PDF URL로 직접 이동
                            await self.page.goto(pdf_link, timeout=60000)

                        download = await download_info.value
                        suggested_filename = download.suggested_filename
                        if suggested_filename:
                            ext = Path(suggested_filename).suffix
                            if ext:
                                filepath = filepath.with_suffix(ext)

                        await download.save_as(str(filepath))

                        # 파일 크기 검증
                        if filepath.exists() and filepath.stat().st_size > 0:
                            print(f"[OK] 실제 파일 저장됨: {filepath}")
                            return str(filepath)
                        else:
                            print(f"[WARN] 파일이 비어있음")
                            if filepath.exists():
                                filepath.unlink()
                    except Exception as nav_err:
                        print(f"[WARN] navigate 다운로드 실패: {nav_err}")

                    # 방법 3: Playwright API Request로 다운로드 (브라우저 세션 쿠키 사용)
                    print(f"  [LINK] Playwright API로 다운로드 시도...")
                    try:
                        api_request = await self.context.request.get(pdf_link)
                        if api_request.ok:
                            content = await api_request.body()
                            content_type = api_request.headers.get('content-type', '')

                            if len(content) > 1000 and ('application/pdf' in content_type or content[:4] == b'%PDF'):
                                if filepath.suffix != '.pdf':
                                    filepath = filepath.with_suffix('.pdf')
                                with open(filepath, 'wb') as f:
                                    f.write(content)
                                print(f"[OK] Playwright API로 저장됨: {filepath}")
                                return str(filepath)
                            else:
                                print(f"[WARN] PDF가 아니거나 내용이 없음 (size: {len(content)}, type: {content_type})")
                        else:
                            print(f"[WARN] API 요청 실패: {api_request.status}")
                    except Exception as api_err:
                        print(f"[WARN] API 다운로드 실패: {api_err}")

                except Exception as e:
                    print(f"[WARN] PDF URL 다운로드 실패: {e}")

            # 다운로드 링크를 찾지 못한 경우
            # Playwright API를 사용하여 다운로드 시도 (브라우저 세션 쿠키 자동 사용)
            print(f"[WARN] 다운로드 링크 없음, Playwright API로 직접 다운로드 시도")

            try:
                api_request = await self.context.request.get(url)
                if api_request.ok:
                    content = await api_request.body()
                    content_type = api_request.headers.get('content-type', '')

                    # PDF 여부 확인 (Content-Type 또는 파일 시그니처로)
                    is_pdf = ('application/pdf' in content_type or
                              url.endswith('.pdf') or
                              (len(content) > 4 and content[:4] == b'%PDF'))

                    if is_pdf and len(content) > 1000:
                        if filepath.suffix != '.pdf':
                            filepath = filepath.with_suffix('.pdf')

                        with open(filepath, 'wb') as f:
                            f.write(content)

                        print(f"[OK] Playwright API로 PDF 저장됨: {filepath}")
                        return str(filepath)
                    else:
                        print(f"[INFO] PDF가 아님 (Content-Type: {content_type}, size: {len(content)}), 건너뜀")
                        return None
                else:
                    print(f"[WARN] API 요청 실패: {api_request.status}")

            except Exception as e:
                print(f"[WARN] Playwright API 다운로드 실패: {e}")

            # headless 모드에서만 page.pdf() 시도
            if self.headless:
                try:
                    if filepath.suffix != '.pdf':
                        filepath = filepath.with_suffix('.pdf')
                    await self._dismiss_cookie_banner()
                    await self.page.pdf(path=str(filepath))
                    print(f"[OK] 페이지 캡처 저장됨: {filepath}")
                    return str(filepath)
                except Exception as e:
                    print(f"[WARN] page.pdf() 실패 (headless 모드에서만 작동): {e}")

            print(f"[SKIP] 다운로드 불가: {url}")
            return None

        except Exception as e:
            print(f"[ERROR] 다운로드 실패: {e}")
            return None

    async def _find_download_element(self):
        """페이지에서 클릭 가능한 다운로드 요소 찾기"""
        try:
            # 다운로드 버튼/링크 셀렉터들 (클릭용)
            download_selectors = [
                # PDF 직접 링크
                'a[href$=".pdf"]',
                'a[href*=".pdf?"]',
                'a[href*="/pdf/"]',
                # 다운로드 속성이 있는 링크
                'a[download]',
                # 텍스트 기반
                'a:has-text("Download PDF")',
                'a:has-text("Download Report")',
                'a:has-text("Download")',
                'a:has-text("다운로드")',
                'a:has-text("PDF")',
                # 클래스 기반
                'a[class*="download"]',
                'a[class*="Download"]',
                '.download-link',
                '.download-btn',
                '.pdf-download',
                # 버튼 형태
                'button:has-text("Download")',
                'button:has-text("다운로드")',
                # 아이콘이 있는 링크 (LIMRA 스타일)
                'a[href*="download"]',
                'a[href*="Download"]',
            ]

            for selector in download_selectors:
                try:
                    elements = await self.page.query_selector_all(selector)
                    for element in elements:
                        if await element.is_visible():
                            href = await element.get_attribute('href')
                            # PDF 링크인지 확인
                            if href and ('.pdf' in href.lower() or 'download' in href.lower()):
                                return element
                except:
                    continue

            return None

        except Exception as e:
            return None

    async def _find_pdf_url(self) -> str:
        """페이지에서 PDF URL 찾기"""
        try:
            # JavaScript로 모든 PDF 링크 수집
            pdf_urls = await self.page.evaluate("""
                () => {
                    const urls = [];
                    // 모든 a 태그에서 PDF 링크 찾기
                    document.querySelectorAll('a[href]').forEach(a => {
                        const href = a.href || '';
                        if (href.toLowerCase().includes('.pdf')) {
                            urls.push(href);
                        }
                    });
                    // iframe 내부도 확인
                    document.querySelectorAll('iframe').forEach(iframe => {
                        try {
                            const src = iframe.src || '';
                            if (src.toLowerCase().includes('.pdf')) {
                                urls.push(src);
                            }
                        } catch(e) {}
                    });
                    // embed/object 태그 확인
                    document.querySelectorAll('embed[src], object[data]').forEach(el => {
                        const src = el.src || el.data || '';
                        if (src.toLowerCase().includes('.pdf')) {
                            urls.push(src);
                        }
                    });
                    return urls;
                }
            """)

            # 첫 번째 유효한 PDF URL 반환
            for url in pdf_urls:
                if url and '.pdf' in url.lower():
                    return url

            return None

        except Exception as e:
            return None

    async def download_all_results(self) -> list:
        """검색된 모든 결과 다운로드"""
        print(f"\n[PKG] {len(self.search_results)}개 문서 다운로드 시작...")

        downloaded_files = []

        for i, result in enumerate(self.search_results, 1):
            print(f"\n[{i}/{len(self.search_results)}] {result['title'][:50]}...")

            # 파일명 생성
            safe_title = re.sub(r'[<>:"/\\|?*]', '_', result['title'])[:100]
            extension = '.pdf' if result['type'] == 'PDF' else '.pdf'
            filename = f"{safe_title}{extension}"

            filepath = await self.download_document(result['url'], filename)
            if filepath:
                downloaded_files.append({
                    'title': result['title'],
                    'url': result['url'],
                    'filepath': filepath
                })

            # 요청 간 딜레이
            await asyncio.sleep(2)

        print(f"\n[OK] 총 {len(downloaded_files)}개 파일 다운로드 완료")
        return downloaded_files

    async def save_results_report(self):
        """검색 결과 리포트 저장"""
        report_path = self.download_folder / f"search_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

        report = {
            'search_date': datetime.now().isoformat(),
            'total_results': len(self.search_results),
            'results': self.search_results
        }

        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        print(f"[REPORT] 검색 리포트 저장됨: {report_path}")
        return report_path

    async def ai_smart_search(
        self,
        keyword: str,
        ai_helper=None,
        max_pages: int = 20
    ) -> list:
        """
        AI 기반 스마트 검색 - Research 탭에서 키워드로 검색 → PDF 다운로드 → AI 분석 → Word 저장

        Args:
            keyword: 검색 키워드
            ai_helper: AI 헬퍼 인스턴스 (필수)
            max_pages: 다운로드할 최대 문서 수

        Returns:
            생성된 보고서 정보
        """
        print("\n" + "="*60)
        print(f"[AI SEARCH] Research 탭 검색 및 AI 분석: '{keyword}'")
        print("="*60)

        if not ai_helper:
            print("[ERROR] AI 헬퍼가 필요합니다.")
            return []

        try:
            # 1. LIMRA 검색 기능으로 키워드 검색
            print(f"\n[STEP 1] LIMRA 사이트에서 '{keyword}' 검색 중...")
            docs = await self.search_documents(query=keyword, max_results=max_pages)

            if not docs:
                print(f"[WARN] '{keyword}' 키워드로 검색된 문서가 없습니다.")
                return []

            print(f"[OK] {len(docs)}개 문서 발견")

            # 최대 다운로드 수 제한
            docs_to_download = docs[:max_pages]
            print(f"[INFO] {len(docs_to_download)}개 문서 다운로드 시작...")

            # 2. 검색된 모든 PDF 다운로드
            print(f"\n[STEP 2] PDF 문서 다운로드 중...")
            self.search_results = docs_to_download
            downloaded_files = await self.download_all_results()

            if not downloaded_files:
                print("[ERROR] 다운로드된 파일이 없습니다.")
                return []

            print(f"[OK] {len(downloaded_files)}개 PDF 다운로드 완료")

            # 3. 다운로드한 모든 PDF를 AI가 읽고 분석
            print(f"\n[STEP 3] AI가 PDF 파일들을 분석 중...")

            # PyPDF2로 PDF 텍스트 추출
            try:
                import PyPDF2
            except ImportError:
                print("[ERROR] PyPDF2가 설치되지 않았습니다. pip install PyPDF2")
                return []

            all_pdf_content = []
            for pdf_file in downloaded_files:
                try:
                    pdf_path = pdf_file['filepath']  # filepath 키 사용
                    print(f"  읽는 중: {pdf_file['title'][:50]}...")

                    with open(pdf_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text_content = ""

                        # 모든 페이지 읽기
                        for page_num, page in enumerate(pdf_reader.pages):
                            try:
                                text_content += page.extract_text() + "\n"
                            except:
                                continue

                        if text_content.strip():
                            all_pdf_content.append({
                                'title': pdf_file['title'],
                                'filepath': str(pdf_path),
                                'content': text_content[:10000]  # 첫 10000자만
                            })
                            print(f"    ✓ {len(text_content)} 글자 추출")

                except Exception as e:
                    print(f"    ✗ PDF 읽기 실패: {e}")
                    continue

            if not all_pdf_content:
                print("[ERROR] PDF에서 텍스트를 추출할 수 없습니다.")
                return []

            print(f"[OK] {len(all_pdf_content)}개 PDF 분석 완료")

            # 4. AI가 모든 문서를 읽고 종합 보고서 작성
            print(f"\n[STEP 4] AI가 종합 보고서 작성 중...")

            # 모든 PDF 내용을 하나로 결합
            combined_text = f"검색 키워드: {keyword}\n\n"
            combined_text += f"총 {len(all_pdf_content)}개 문서 분석\n\n"

            for i, doc in enumerate(all_pdf_content, 1):
                combined_text += f"\n{'='*60}\n"
                combined_text += f"문서 {i}: {doc['title']}\n"
                combined_text += f"파일: {doc['filepath']}\n"
                combined_text += f"{'='*60}\n"
                combined_text += doc['content'][:5000] + "\n\n"

            # AI 프롬프트
            prompt = f"""당신은 보험 산업 리서치 분석가입니다. 다음은 LIMRA Research 섹션에서 '{keyword}' 키워드로 검색한 {len(all_pdf_content)}개의 PDF 문서 내용입니다.

{combined_text[:50000]}

위 문서들을 종합 분석하여 다음 형식으로 상세한 한국어 보고서를 작성해주세요:

# {keyword} - LIMRA Research 종합 분석 보고서

## 1. 요약 (Executive Summary)
전체 문서의 핵심 내용을 3-5문단으로 요약

## 2. 주요 발견사항 (Key Findings)
- 각 문서에서 발견된 중요한 인사이트
- 데이터, 통계, 트렌드
- 산업에 대한 시사점

## 3. 상세 분석 (Detailed Analysis)
각 문서별로:
- 문서 제목
- 핵심 내용
- 주요 데이터 포인트
- 분석 및 해석

## 4. 결론 및 권장사항 (Conclusions & Recommendations)
- 전체 문서를 통해 얻은 통찰
- 비즈니스 의사결정을 위한 권장사항
- 향후 고려사항

## 5. 참고 문헌 (References)
분석한 모든 문서 목록

보고서는 전문적이고 상세하게 작성해주세요."""

            # Groq AI 호출
            from groq import Groq
            client = Groq(api_key=ai_helper.api_key)

            response = client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=ai_helper.model,
                temperature=0.3,
                max_tokens=6000
            )

            report = response.choices[0].message.content
            print(f"[OK] AI 보고서 생성 완료 ({len(report)} 글자)")

            # 5. Word 문서로 저장
            print(f"\n[STEP 5] Word 문서(.docx)로 저장 중...")

            try:
                from docx import Document
                from docx.shared import Pt, RGBColor, Inches
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                # Word 문서 생성
                doc = Document()

                # 제목 추가
                title = doc.add_heading(f'{keyword} - LIMRA Research 종합 분석 보고서', 0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # 메타데이터
                meta = doc.add_paragraph()
                meta.add_run(f'생성일: {datetime.now().strftime("%Y년 %m월 %d일")}\n')
                meta.add_run(f'분석 문서 수: {len(all_pdf_content)}개\n')
                meta.add_run(f'키워드: {keyword}\n')
                meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                doc.add_paragraph('_' * 50)

                # 보고서 내용 추가
                for line in report.split('\n'):
                    if line.startswith('# '):
                        # 대제목
                        doc.add_heading(line.replace('# ', ''), 1)
                    elif line.startswith('## '):
                        # 중제목
                        doc.add_heading(line.replace('## ', ''), 2)
                    elif line.startswith('### '):
                        # 소제목
                        doc.add_heading(line.replace('### ', ''), 3)
                    elif line.strip().startswith('- ') or line.strip().startswith('* '):
                        # 리스트
                        doc.add_paragraph(line.strip()[2:], style='List Bullet')
                    elif line.strip():
                        # 일반 텍스트
                        doc.add_paragraph(line)

                # 파일명 생성
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"LIMRA_AI_Report_{keyword.replace(' ', '_')}_{timestamp}.docx"
                filepath = self.download_folder / filename

                doc.save(str(filepath))
                print(f"[OK] Word 문서 저장 완료: {filename}")

                return [{
                    'title': f'{keyword} AI 종합 분석 보고서',
                    'url': 'AI Generated',
                    'type': 'AI Report (Word)',
                    'file': filename,
                    'pdf_count': len(all_pdf_content),
                    'downloaded_files': [f['filepath'] for f in downloaded_files]
                }]

            except ImportError:
                print("[WARN] python-docx가 설치되지 않았습니다. 텍스트 파일로 저장합니다.")

                # Word 라이브러리 없으면 TXT로 저장
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"LIMRA_AI_Report_{keyword.replace(' ', '_')}_{timestamp}.txt"
                filepath = self.download_folder / filename

                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(f"{keyword} - LIMRA Research 종합 분석 보고서\n")
                    f.write(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일')}\n")
                    f.write(f"분석 문서 수: {len(all_pdf_content)}개\n\n")
                    f.write("="*60 + "\n\n")
                    f.write(report)

                print(f"[OK] 텍스트 파일 저장 완료: {filename}")

                return [{
                    'title': f'{keyword} AI 종합 분석 보고서',
                    'url': 'AI Generated',
                    'type': 'AI Report (TXT)',
                    'file': filename,
                    'pdf_count': len(all_pdf_content),
                    'downloaded_files': [f['filepath'] for f in downloaded_files]
                }]

        except Exception as e:
            print(f"[ERROR] AI 스마트 검색 실패: {e}")
            import traceback
            traceback.print_exc()
            return []

    async def close(self):
        """브라우저 종료"""
        if self.browser:
            await self.browser.close()
            print("[CLOSE] 브라우저 종료됨")


async def main():
    """메인 실행 함수"""
    # 설정
    EMAIL = "plopa1535@kyobo.com"
    PASSWORD = "Kyobo1234!@#$"
    DOWNLOAD_FOLDER = "./limra_downloads"

    # 에이전트 생성
    agent = LimraSearchAgent(
        email=EMAIL,
        password=PASSWORD,
        download_folder=DOWNLOAD_FOLDER,
        headless=False  # 브라우저 창 표시 (디버깅용)
    )

    try:
        # 초기화
        await agent.initialize()

        # 로그인
        login_success = await agent.login()

        if login_success:
            print("\n" + "="*50)
            print("로그인 성공! 이제 검색을 수행할 수 있습니다.")
            print("="*50)

            # 메뉴 선택
            print("\n[MENU] 작업을 선택하세요:")
            print("1. 키워드 검색")
            print("2. Research 섹션 탐색 (주제/연도 필터링)")
            print("3. Research 섹션 전체 탐색")

            menu_choice = input("\n선택 (1/2/3): ").strip()

            if menu_choice == '1':
                # 키워드 검색
                search_query = input("\n[SEARCH] 검색어를 입력하세요: ").strip()

                if search_query:
                    results = await agent.search_documents(search_query, max_results=20)

                    print("\n[LIST] 검색 결과:")
                    print("-" * 50)
                    for i, result in enumerate(results, 1):
                        print(f"{i}. [{result['type']}] {result['title'][:60]}")
                        print(f"   URL: {result['url']}")
                        print()

                    if results:
                        download_choice = input("\n모든 결과를 다운로드하시겠습니까? (y/n): ").strip().lower()
                        if download_choice == 'y':
                            downloaded = await agent.download_all_results()
                            print(f"\n[OK] {len(downloaded)}개 파일이 {DOWNLOAD_FOLDER}에 저장되었습니다.")

                    await agent.save_results_report()

            elif menu_choice == '2':
                # Research 섹션 탐색 (주제/연도 필터링)
                print("\n" + "="*50)
                print("[DOC] Research 섹션 필터링 검색")
                print("="*50)

                # 주제 키워드 입력
                print("\n[SEARCH] 주제 키워드를 입력하세요")
                print("   (쉼표로 구분, 예: Recruiting, Retention, Agent)")
                print("   (전체 검색은 빈칸으로 Enter)")
                keywords_input = input("   키워드: ").strip()

                keywords = None
                if keywords_input:
                    keywords = [kw.strip() for kw in keywords_input.split(',') if kw.strip()]

                # 연도 범위 입력
                print("\n[DATE] 연도 범위를 입력하세요")
                print("   (전체 기간은 빈칸으로 Enter)")

                start_year_input = input("   시작 연도 (예: 2023): ").strip()
                end_year_input = input("   종료 연도 (예: 2024): ").strip()

                start_year = int(start_year_input) if start_year_input.isdigit() else None
                end_year = int(end_year_input) if end_year_input.isdigit() else None

                # 자동 다운로드 여부
                auto_download = input("\n[DL] 필터링된 문서를 자동 다운로드할까요? (y/n): ").strip().lower() == 'y'

                # 필터링 검색 실행
                filtered_docs = await agent.browse_research_with_filter(
                    keywords=keywords,
                    start_year=start_year,
                    end_year=end_year,
                    auto_download=auto_download
                )

                # 결과 저장
                if filtered_docs:
                    filter_desc = []
                    if keywords:
                        filter_desc.append(f"keywords_{'_'.join(keywords)}")
                    if start_year or end_year:
                        filter_desc.append(f"{start_year or 'any'}-{end_year or 'any'}")

                    report_name = f"filtered_research_{'_'.join(filter_desc)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                    report_path = agent.download_folder / report_name

                    with open(report_path, 'w', encoding='utf-8') as f:
                        json.dump({
                            'filter_date': datetime.now().isoformat(),
                            'keywords': keywords,
                            'start_year': start_year,
                            'end_year': end_year,
                            'total_documents': len(filtered_docs),
                            'documents': filtered_docs
                        }, f, ensure_ascii=False, indent=2)
                    print(f"\n[FILE] 필터링 결과 저장됨: {report_path}")

                    # 수동 다운로드 옵션
                    if not auto_download and filtered_docs:
                        download_choice = input("\n결과를 다운로드하시겠습니까? (y/n): ").strip().lower()
                        if download_choice == 'y':
                            max_download = input(f"다운로드할 최대 문서 수 (기본: 전체 {len(filtered_docs)}개): ").strip()
                            max_download = int(max_download) if max_download.isdigit() else len(filtered_docs)

                            agent.search_results = filtered_docs[:max_download]
                            downloaded = await agent.download_all_results()
                            print(f"\n[OK] {len(downloaded)}개 파일이 {DOWNLOAD_FOLDER}에 저장되었습니다.")

            elif menu_choice == '3':
                # 연구 섹션 전체 탐색
                docs = await agent.browse_research_section()
                print(f"\n[DOC] {len(docs)}개의 연구 문서를 발견했습니다.")

                research_report_path = agent.download_folder / f"research_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                with open(research_report_path, 'w', encoding='utf-8') as f:
                    json.dump({
                        'scan_date': datetime.now().isoformat(),
                        'total_documents': len(docs),
                        'documents': docs
                    }, f, ensure_ascii=False, indent=2)
                print(f"[FILE] 연구 문서 목록 저장됨: {research_report_path}")

                if docs:
                    download_research = input("\n연구 문서를 다운로드하시겠습니까? (y/n): ").strip().lower()
                    if download_research == 'y':
                        max_download = input(f"다운로드할 최대 문서 수 (기본: 20, 전체: {len(docs)}): ").strip()
                        max_download = int(max_download) if max_download.isdigit() else 20

                        agent.search_results = docs[:max_download]
                        downloaded = await agent.download_all_results()
                        print(f"\n[OK] {len(downloaded)}개 연구 문서가 {DOWNLOAD_FOLDER}에 저장되었습니다.")

        else:
            print("\n[ERROR] 로그인에 실패했습니다. 자격 증명을 확인해주세요.")
            print("   다운로드 폴더에서 디버그 스크린샷을 확인하세요.")

    except KeyboardInterrupt:
        print("\n\n[WARN] 사용자에 의해 중단됨")

    except Exception as e:
        print(f"\n[ERROR] 오류 발생: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    # Windows asyncio 경고 숨기기
    import warnings
    warnings.filterwarnings("ignore", category=ResourceWarning)

    # Windows에서 ProactorEventLoop 관련 경고 방지
    import sys
    if sys.platform == 'win32':
        asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

    asyncio.run(main())
